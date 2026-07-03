#!/usr/bin/env python3
"""
研报生成器 - 八字五行+9因子选股
命主: 1994年2月20日戌时男
八字: 甲戌 丙寅 丁丑 庚戌 → 土极旺, 用神金水
"""
import sys, os, json, io, warnings, time
warnings.filterwarnings("ignore")
sys.path.insert(0, os.path.expanduser("~/my-codespace"))

from datetime import datetime
import numpy as np
import pandas as pd
from pathlib import Path

# ── 新引擎 ──
from technicals import full_analysis, TechnicalSummary
from multifactor import full_factor_score, FactorScores

# ── 八字五行映射 ──
# 股票代码河图数: 1壬水 2乙木 3丙火 4丁火 5戊土 6癸水 7庚金 8辛金 9壬水 0己土
HETU = {"1":"水","2":"木","3":"火","4":"火","5":"土","6":"水","7":"金","8":"金","9":"水","0":"土"}

# 行业五行映射
INDUSTRY_WUXING = {
    "证券":"金","银行":"金","保险":"金","有色/黄金":"金","有色金属":"金",
    "钢铁":"金","汽车":"金","汽车零部件":"金","金融IT":"金",
    "白酒":"水","航运":"水","航空":"水","传媒":"水","互联网":"水",
    "AI":"水","AI应用":"水","AI金融":"水","AI芯片":"水",
    "半导体":"水","芯片":"水","消费电子":"水","通信":"水",
    "游戏":"水","软件/AI":"水","电商":"水",
    "新能源车":"火","动力电池":"火","光伏":"火","光模块":"火",
    "锂矿":"火","锂电池":"火","稀土":"火",
    "创新药":"土","CXO":"土","医药":"土","医疗":"土","医疗器械":"土",
    "中药":"土","生物医药":"土","疫苗":"土",
    "银行":"金","保险":"金","证券":"金",
    "光伏逆变器":"火","工控":"木","安防":"土","面板":"土",
    "PCB":"金","AI制造":"火","设备":"金",
    "AI":"水","AI应用":"水",
    "煤炭":"土","石油":"金","电力":"火","建材":"土",
    "家电":"火","物流":"土","铁路":"土","交易所":"金",
    "本地生活":"水","餐饮":"土","乳业":"土","调味品":"土",
    "酒":"水","新能源":"火","免税":"水"
}

def code_wuxing(code):
    """从股票代码取五行"""
    # 取后几位数字
    digits = [c for c in code if c.isdigit()]
    if not digits:
        return "土", 0
    # 尾数权重最大, 取尾2位
    last = digits[-1]
    second_last = digits[-2] if len(digits) >= 2 else "5"
    elem_last = HETU.get(last, "土")
    elem_second = HETU.get(second_last, "土")
    score = 0
    if elem_last in ("金","水"):
        score += 40
    if elem_second in ("金","水"):
        score += 20
    return elem_last, score

def industry_wuxing(sector):
    """行业五行"""
    elem = INDUSTRY_WUXING.get(sector, "土")
    return elem

def name_wuxing(name):
    """名称笔画五行 (简化: 看关键字)"""
    jin_chars = {"金","银","铜","铁","钢","鑫","锋","剑","瑞","锐","铭","锦","钊"}
    shui_chars = {"水","江","河","湖","海","泉","雨","雪","冰","润","泽","源","浩","潇","澜","泰","鸿"}
    huochai = {"火","光","电","炎","灿","煌","晶","明","阳","辉"}
    tuchai = {"土","山","石","岩","峰","地","城","安","宇","奥"}
    mujian = {"木","林","森","植","东"}
    
    score = 0
    for c in name:
        if c in jin_chars: score += 15
        elif c in shui_chars: score += 15
        elif c in huochai: score -= 5
        elif c in tuchai: score -= 5
    return min(40, max(-20, score))

def bazi_score(code, name, sector):
    """综合八字五行评分 0~100"""
    ce, cs = code_wuxing(code)
    ie = industry_wuxing(sector)
    ns = name_wuxing(name)
    
    score = 50  # 基础分
    
    # 代码五行 - 用神金+40, 水+30
    if ce == "金": score += cs * 1.0  # 40
    elif ce == "水": score += cs * 0.8  # 32
    
    # 行业五行
    if ie == "金": score += 20
    elif ie == "水": score += 15
    elif ie == "火": score -= 10
    elif ie == "土": score -= 8
    
    # 名称
    score += ns
    
    return max(0, min(100, score)), ce, ie, ns

def bazi_detail(code, name, sector):
    """五行详细分解"""
    score, ce, ie, ns = bazi_score(code, name, sector)
    parts = []
    if ce in ("金","水"): parts.append(f"尾数{ce}✅")
    else: parts.append(f"尾数{ce}")
    if ie in ("金","水"): parts.append(f"行业{ie}✅")
    elif ie == "火": parts.append(f"行业{ie}⚠忌")
    elif ie == "土": parts.append(f"行业{ie}⚠忌")
    else: parts.append(f"行业{ie}")
    if ns > 0: parts.append(f"名称+{ns:.0f}✅")
    elif ns < 0: parts.append(f"名称{ns:.0f}⚠")
    return score, " ".join(parts)

# ── 主流程 ──
def generate_report():
    out_dir = Path.home() / "Desktop"
    out_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    
    # 加载股票池
    POOL = [
        # AI/科技
        ("002230.SZ","科大讯飞","AI"),("300418.SZ","昆仑万维","AI应用"),
        ("300033.SZ","同花顺","AI金融"),("688111.SS","金山办公","软件/AI"),
        ("688256.SS","寒武纪","AI芯片"),("002049.SZ","紫光国微","半导体"),
        ("603986.SS","兆易创新","半导体"),("688981.SS","中芯国际","半导体"),
        ("002371.SZ","北方华创","半导体设备"),("002475.SZ","立讯精密","消费电子"),
        # 新能源/汽车
        ("002594.SZ","比亚迪","新能源车"),("300750.SZ","宁德时代","动力电池"),
        ("002460.SZ","赣锋锂业","锂矿"),("601012.SS","隆基绿能","光伏"),
        # 金融
        ("600036.SS","招商银行","银行"),("601318.SS","中国平安","保险"),
        ("600030.SS","中信证券","证券"),("300059.SZ","东方财富","证券"),
        ("601899.SS","紫金矿业","有色/黄金"),("600111.SS","北方稀土","稀土"),
        # 消费/酒
        ("600519.SS","贵州茅台","白酒"),("000858.SZ","五粮液","白酒"),
        ("000568.SZ","泸州老窖","白酒"),("600809.SS","山西汾酒","白酒"),
        # 医药
        ("600276.SS","恒瑞医药","创新药"),("603259.SS","药明康德","CXO"),
        ("300347.SZ","泰格医药","CXO"),("688180.SS","君实生物","创新药"),
        # 港股
        ("0700.HK","腾讯控股","互联网"),("9988.HK","阿里巴巴","电商"),
        ("9999.HK","网易","游戏"),("3690.HK","美团","本地生活"),
        ("9618.HK","京东","电商"),("9888.HK","百度","AI"),
        ("1810.HK","小米集团","消费电子"),("0388.HK","港交所","交易所"),
        ("0941.HK","中国移动","通信"),("0883.HK","中国海油","石油"),
        # 中特估
        ("601857.SS","中国石油","石油"),("600900.SS","长江电力","电力"),
        ("601728.SS","中国电信","通信"),("600941.SS","中国移动","通信"),
        ("600585.SS","海螺水泥","建材"),("601919.SS","中远海控","航运"),
        # 更多
        ("300896.SZ","爱美客","医美"),("688008.SS","澜起科技","半导体"),
        ("002920.SZ","德赛西威","汽车电子"),("000725.SZ","京东方A","面板"),
        ("002415.SZ","海康威视","安防"),("600570.SS","恒生电子","金融IT"),
        ("603501.SS","韦尔股份","芯片"),("300308.SZ","中际旭创","光模块"),
        ("300502.SZ","新易盛","光模块"),("002555.SZ","三七互娱","游戏"),
        ("300015.SZ","爱尔眼科","医疗"),("300122.SZ","智飞生物","疫苗"),
        ("000651.SZ","格力电器","家电"),("000333.SZ","美的集团","家电"),
        ("002352.SZ","顺丰控股","物流"),("600104.SS","上汽集团","汽车"),
        ("000625.SZ","长安汽车","汽车"),("601633.SS","长城汽车","汽车"),
        ("002466.SZ","天齐锂业","锂矿"),("300413.SZ","芒果超媒","传媒"),
        ("002027.SZ","分众传媒","传媒"),("300773.SZ","拉卡拉","支付"),
        ("000997.SZ","新大陆","数字支付"),("002152.SZ","广电运通","数币"),
        ("600570.SS","恒生电子","金融IT"),("603288.SS","海天味业","调味品"),
        ("600887.SS","伊利股份","乳业"),("600809.SS","山西汾酒","白酒"),
    ]
    
    print("📡 获取数据...")
    from stock_cache import StockCache
    cache = StockCache()
    symbols = [p[0] for p in POOL]
    raw = cache.get_data(symbols)
    if isinstance(raw, tuple): data_dict = raw[0]
    else: data_dict = raw
    
    results = {}
    for k, df in data_dict.items():
        if df is not None and not df.empty:
            df.columns = [c.lower() for c in df.columns]
            results[k] = df
    
    print(f"📊 评分中 ({len(results)}只)...")
    
    scored = []
    for code, name, sector in POOL:
        df = results.get(code)
        if df is None or df.empty: continue
        
        # 八字五行分
        bz_score, bz_parts = bazi_detail(code, name, sector)
        
        # 技术面
        ts = full_analysis(code, df)
        from technicals import score_technical
        tech_score = score_technical(ts)
        
        # 9因子总分
        fs, ts2 = full_factor_score(code, df)
        
        # 综合推荐分: 五行30% + 技术20% + 因子50%
        final = bz_score * 0.30 + tech_score * 0.20 + fs.total * 0.50
        
        scored.append({
            "code": code, "name": name, "sector": sector,
            "bazi": round(bz_score, 1), "bazi_detail": bz_parts,
            "tech": round(tech_score, 1), "factor": round(fs.total, 1),
            "final": round(final, 1),
            "rsi": round(ts.rsi.value, 1) if ts.rsi else 0,
            "macd": ts.macd.signal_type if ts.macd else "?",
            "trend": ts.trend, "atr": round(ts.atr_pct, 1),
            "volume_ratio": round(ts.volume.volume_ratio, 1) if ts.volume else 1.0
        })
    
    # 排序: 综合分 + 五行优先
    scored.sort(key=lambda x: (x['final'], x['bazi']), reverse=True)
    
    return scored, timestamp, out_dir

def print_report(scored, timestamp, out_dir):
    print(f"\n{'='*80}")
    print(f"  📜 八字五行选股研报")
    print(f"  命主: 1994.02.20 戌时  |  用神: 金、水  |  忌: 火、土")
    print(f"  生成: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print(f"{'='*80}")
    print(f"\n{'─'*80}")
    print(f"  🔥 Top 20 金水推荐 (按综合分排序)")
    print(f"{'─'*80}")
    print(f"  {'排名':<4s} {'代码':<12s} {'名称':<10s} {'行业':<8s} {'综合':>5s} {'五行':>5s} {'技术':>5s} {'因子':>5s} {'RSI':>5s} {'趋势':>6s}")
    print(f"  {'─'*4} {'─'*12} {'─'*10} {'─'*8} {'─'*5} {'─'*5} {'─'*5} {'─'*5} {'─'*5} {'─'*6}")
    
    for i, s in enumerate(scored[:20], 1):
        trend_i = {"uptrend":"📈","weak_uptrend":"↗","sideways":"➡","weak_downtrend":"↘","downtrend":"📉"}.get(s['trend'],"➡")
        print(f"  {i:<4d} {s['code']:<12s} {s['name']:<10s} {s['sector']:<8s} {s['final']:>5.1f} {s['bazi']:>5.1f} {s['tech']:>5.1f} {s['factor']:>5.1f} {s['rsi']:>5.1f} {trend_i:>6s}")
    
    # Top 5 详细
    print(f"\n{'═'*80}")
    print(f"  📝 Top 5 深度分析")
    print(f"{'═'*80}")
    
    for i, s in enumerate(scored[:5], 1):
        print(f"\n  ╔══ #{i} {s['name']} ({s['code']}) ── {s['sector']}")
        print(f"  ║  综合分:{s['final']}  八字:{s['bazi']}  技术:{s['tech']}  因子:{s['factor']}")
        print(f"  ║  {s['bazi_detail']}")
        print(f"  ║  RSI:{s['rsi']}  MACD:{s['macd']}  趋势:{s['trend']}  ATR:{s['atr']}%  量比:{s['volume_ratio']}")
        print(f"  ╚══")
    
    # 操作建议
    print(f"\n{'═'*80}")
    print(f"  操作建议 (6万→8万计划)")
    print(f"{'═'*80}")
    
    # 选金水最好的前3
    top3 = sorted(scored, key=lambda x: x['bazi'], reverse=True)[:3]
    print(f"\n  匹配八字最旺的3只:")
    for i, s in enumerate(top3[:3], 1):
        alloc = "2万" if i < 3 else "2万"
        target = f"+{(33/len(top3[:3])):.0f}%" 
        print(f"  {i}. {s['name']} ({s['code']})  八字{s['bazi']}分  分配{alloc}  目标+{int(33/len(top3[:3]))}%")
    
    print(f"\n  ⚠ 免责: 股市有风险, 八字仅供参考, 不构成投资建议")
    return scored[:20]

def generate_docx(scored, timestamp, out_dir):
    """生成标准研报 .docx"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # ── 样式设置 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.75
    
    # ── 标题 ──
    title = doc.add_heading('', 0)
    run = title.add_run('股票投资研究报告')
    run.font.size = Pt(22)
    run.font.name = '方正小标宋简体'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'—— 基于八字命理与多因子量化选股模型')
    run.font.size = Pt(14)
    run.font.name = '楷体'
    
    # ── 基本信息表 ──
    doc.add_paragraph()
    info = doc.add_table(rows=5, cols=2)
    info.style = 'Table Grid'
    cells = [
        ('报告日期', datetime.now().strftime('%Y年%m月%d日')),
        ('命主信息', '1994年2月20日戌时（甲戌 丙寅 丁丑 庚戌）'),
        ('五行格局', '日主丁火，土极旺，缺金少水'),
        ('用神/忌神', '用神：金、水　｜　忌神：火、土'),
        ('选股策略', '金水行业优先 + 河图数尾1/6/9 + 9因子评分'),
    ]
    for i, (k, v) in enumerate(cells):
        info.cell(i, 0).text = k
        info.cell(i, 1).text = v
        for cell in [info.cell(i, 0), info.cell(i, 1)]:
            for p in cell.paragraphs:
                p.style.font.size = Pt(12)
    
    # ── Top 10 推荐 ──
    doc.add_heading('一、金水优选股票池', level=1)
    
    top10 = scored[:10]
    table = doc.add_table(rows=len(top10)+1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['排名', '代码', '名称', '行业', '综合分', '八字分', 'RSI', '操作建议']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True
    
    for i, s in enumerate(top10):
        bazi_mark = "⭐" if s['bazi'] >= 60 else ""
        if s['rsi'] < 30: action = "超跌布局"
        elif s['rsi'] > 70: action = "注意回调"
        elif s['trend'] in ('uptrend','weak_uptrend'): action = "顺势持有"
        elif s['trend'] in ('sideways',): action = "等待突破"
        else: action = "观望"
        
        row_data = [str(i+1), s['code'], s['name'], s['sector'],
                    f"{s['final']:.0f}", f"{s['bazi']:.0f}{bazi_mark}",
                    f"{s['rsi']:.0f}", action]
        for j, v in enumerate(row_data):
            table.cell(i+1, j).text = v
    
    # ── 深度分析 ──
    doc.add_heading('二、核心标的深度分析', level=1)
    
    for i, s in enumerate(top10[:5]):
        doc.add_heading(f'2.{i+1} {s["name"]}（{s["code"]}）', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run(f'【八字匹配】{s["bazi_detail"]}　综合评分{s["bazi"]}分')
        run.bold = True
        
        doc.add_paragraph(
            f'技术面：RSI {s["rsi"]}，MACD {s["macd"]}，'
            f'趋势{s["trend"]}，ATR {s["atr"]}%，量比{s["volume_ratio"]}。'
            f'综合分{s["final"]}（五行30%+技术20%+因子50%）。',
            style='List Bullet'
        )
        
        # 建议
        if s['rsi'] < 30:
            adv = '超卖区域，建议分批建仓。'
        elif s['rsi'] > 70:
            adv = '超买区域，建议等待回调。'
        elif s['trend'] in ('uptrend', 'weak_uptrend'):
            adv = '上升趋势，顺势持有，止损设-7%。'
        else:
            adv = '横盘整理，等待突破信号。'
        
        p = doc.add_paragraph(f'操作建议：{adv}')
        run = p.runs[0]
        run.bold = True
        
        # 9因子分解
        doc.add_paragraph(
            f'9因子：价值{s["tech"]:.0f} 技术{s["tech"]:.0f} '
            f'因子总分{s["factor"]:.0f}/100',
            style='List Bullet'
        )
    
    # ── 操作计划 ──
    doc.add_heading('三、操作计划（6万→8万）', level=1)
    
    top3_gold = sorted(scored, key=lambda x: x['bazi'], reverse=True)[:3]
    
    plan_table = doc.add_table(rows=4, cols=5)
    plan_table.style = 'Table Grid'
    plan_headers = ['标的', '代码', '分配', '目标收益', '止损线']
    for j, h in enumerate(plan_headers):
        plan_table.cell(0, j).text = h
        for p in plan_table.cell(0, j).paragraphs:
            for r in p.runs: r.bold = True
    
    allocs = [3, 2, 1]
    targets = [30, 25, 20]
    stops = [7, 6, 5]
    for i, s in enumerate(top3_gold):
        plan_table.cell(i+1, 0).text = s['name']
        plan_table.cell(i+1, 1).text = s['code']
        plan_table.cell(i+1, 2).text = f"{allocs[i]}万"
        plan_table.cell(i+1, 3).text = f"+{targets[i]}%"
        plan_table.cell(i+1, 4).text = f"-{stops[i]}%"
    
    total_p = doc.add_paragraph(f'\n总投入：6万　预期收益：+33%（约8万）　风险敞口：-7%')
    
    # ── 免责 ──
    doc.add_heading('免责声明', level=3)
    doc.add_paragraph(
        '本报告基于公开数据与八字命理模型生成，仅供参考，不构成投资建议。'
        '股市有风险，投资需谨慎。过往表现不代表未来收益。'
    )
    
    # ── 保存 ──
    filename = f'股票研报_{timestamp}.docx'
    path = out_dir / filename
    doc.save(str(path))
    print(f"\n💾 研报已保存: {path}")
    return path

if __name__ == "__main__":
    print("八字五行选股研报生成器 v1.0")
    print("="*60)
    
    scored, timestamp, out_dir = generate_report()
    print_report(scored, timestamp, out_dir)
    path = generate_docx(scored, timestamp, out_dir)
    print(f"\n✅ 完成! 研报在桌面: {path}")
