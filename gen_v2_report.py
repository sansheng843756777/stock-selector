#!/usr/bin/env python3
"""更新研报 - 6万→8万计划版"""
import sys, os, warnings
warnings.filterwarnings("ignore")
from datetime import datetime
from pathlib import Path

OUT = Path.home() / "Desktop"

# ── 核心金水数据 ──
STOCKS = [
    # (名称, 代码, 行业, 行业五行, 尾数, 尾数五行, RSI, ATR%, 趋势, 量比, 5日%, MACD)
    ("昆仑万维","300418.SZ","AI应用","水",8,"金", 70.3,6.4,"📈上升",2.6,+11.6,"🟢"),
    ("东方财富","300059.SZ","证券","金",9,"水", 67.7,5.5,"📈上升",1.1,-1.8,"🟢"),
    ("金山办公","688111.SS","软件/AI","水",1,"水", 54.4,4.4,"➡横盘",1.4,+1.7,"🟢"),
    ("寒武纪","688256.SS","AI芯片","水",6,"水", 58.1,7.8,"➡横盘",1.0,-8.8,"🔴"),
    ("中信证券","600030.SS","证券","金",0,"土", 64.9,3.9,"📈上升",1.1,-1.2,"🟢"),
    ("紫金矿业","601899.SS","有色/黄金","金",9,"水", 47.8,5.8,"➡横盘",1.4,+3.0,"🔴"),
    ("同花顺","300033.SZ","AI金融","水",3,"火", 69.4,6.8,"📈上升",1.1,-4.4,"🟢"),
    ("澜起科技","688008.SS","半导体","水",8,"金", 59.2,8.5,"➡横盘",1.3,-3.6,"🟢"),
    ("恒生电子","600570.SS","金融IT","金",0,"土", 51.7,4.8,"➡横盘",1.1,-5.6,"🟢"),
    ("分众传媒","002027.SZ","传媒","水",7,"金", 29.9,3.5,"➡横盘",1.3,+1.9,"🟢"),
]

def calc_score(name,code,sector,ie,ld,ce,rsi,atr,trend,vratio,ret5,macd):
    """五行+技术综合评分"""
    wx = 50
    if ie == "金": wx += 20
    elif ie == "水": wx += 15
    elif ie in ("火","土"): wx -= 10
    if ce == "水": wx += 30
    elif ce == "金": wx += 25
    
    tech = 50
    if "上升" in trend: tech += 20
    if "横盘" in trend: tech += 5
    if macd == "🟢": tech += 10
    if vratio > 2: tech += 15
    elif vratio > 1.3: tech += 8
    if 30 <= rsi <= 60: tech += 8
    if rsi < 30: tech += 12  # 超卖反弹
    tech += min(15, atr * 2)
    if ret5 > 8: tech += 5
    
    return round(wx * 0.35 + tech * 0.65)

scored = []
for s in STOCKS:
    name,code,sector,ie,ld,ce,rsi,atr,trend,vratio,ret5,macd = s
    total = calc_score(name,code,sector,ie,ld,ce,rsi,atr,trend,vratio,ret5,macd)
    scored.append(s + (total,))
scored.sort(key=lambda x: x[-1], reverse=True)

# ── 生成 .docx ──
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.75

# 标题
title = doc.add_heading('', 0)
r = title.add_run('股票投资研究报告（一周6万→8万计划）')
r.font.size = Pt(22)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(f'——八字命理+多因子量化  {datetime.now().strftime("%Y-%m-%d %H:%M")}')
r.font.size = Pt(14)

# 命主
doc.add_heading('命主信息', level=1)
for k,v in [("出生","1994年2月20日戌时"),("八字","甲戌 丙寅 丁丑 庚戌"),
            ("五行","土极旺，缺金少水"),("用神","金、水（证券/有色/AI/软件/白酒/半导体）"),
            ("忌神","火、土（新能源/光伏/医药/煤炭）")]:
    p = doc.add_paragraph()
    r = p.add_run(f'{k}：'); r.bold = True
    p.add_run(v)

# 操作目标
doc.add_heading('一、操作目标：6万→8万（一周+33%）', level=1)
p = doc.add_paragraph('核心策略：重仓高位量金水票，利用20%涨跌幅快速获利。')
p.runs[0].bold = True

# 金水池
doc.add_heading('二、金水优选（综合排名）', level=1)
t = doc.add_table(rows=len(scored)+1, cols=9)
t.style = 'Table Grid'
for j,h in enumerate(['排名','名称','代码','行业','五行','RSI','ATR%','趋势','综合']):
    c = t.cell(0,j); c.text = h
    for p in c.paragraphs: 
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in p.runs: rr.bold = True

for i,s in enumerate(scored):
    n,c,se,ie,ld,ce,rsi,atr,trend,vr,rt5,macd,total = s
    wx_mark = f"{se}({ie})"
    for j,v in enumerate([str(i+1), n, c, wx_mark, f"{ld}({ce})", f"{rsi:.0f}", f"{atr:.1f}%", trend, str(total)]):
        t.cell(i+1,j).text = v

# Top3 深度
doc.add_heading('三、核心推荐', level=1)

details = [
    [
        "量比2.6超级放量，5日+11.6%已启动加速",
        "AI概念最强势主线，昆仑万维AI应用龙头",
        "RSI 70.3偏热但量能充沛，趋势延续概率大",
        "尾数8金+行业AI水 → 金水双全",
        "创业板20%涨跌幅，一周内冲击+33%空间足",
        "⚡操作：周一若低开/平开直接进，高开+5%以上等回调"
    ],
    [
        "证券龙头+金水双全（行业金+尾数9水）",
        "趋势向上+MACD多头+上升通道",
        "资金稳健流入，牛市旗手",
        "ATR 5.5%弹性好，适合做底仓",
        "⚡操作：开盘直接入，不追高"
    ],
    [
        "代码688111 → 三壬水（水水水）金水最旺",
        "名称含【金】字 → 五行完美",
        "RSI 54.4中性，量比1.4放量",
        "科创板20%涨跌幅，适合防守仓",
        "⚡操作：昆仑万维止盈后切换，或作为备用"
    ],
]

for idx, s in enumerate(scored[:3]):
    n,c,se,ie,ld,ce,rsi,atr,trend,vr,rt5,macd,total = s
    doc.add_heading(f'{idx+1}. {n}（{c}）', level=2)
    p = doc.add_paragraph()
    r = p.add_run(f'⭐ 综合评分{total}分 | RSI:{rsi:.0f} | ATR:{atr:.1f}% | {trend} | 5日{rt5:+.1f}%')
    r.bold = True
    p = doc.add_paragraph(f'【五行】行业{se}属{ie}✅ | 尾号{ld}属{ce}✅ | MACD:{macd}')
    for d in details[idx]:
        doc.add_paragraph(d, style='List Bullet')

# 操作计划
doc.add_heading('四、6万→8万执行计划', level=1)
pt = doc.add_table(rows=4, cols=7)
pt.style = 'Table Grid'
for j,h in enumerate(['标的','代码','分配','入场','目标','止损','收益']):
    c = pt.cell(0,j); c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in p.runs: rr.bold = True

plans = [
    ("昆仑万维","300418.SZ","4万","RSI回调至60~65或周一开盘平开","+35%","-7%","+1.4万"),
    ("东方财富","300059.SZ","2万","开盘直接入","+30%","-5%","+0.6万"),
    ("金山办公","688111.SS","备用","昆仑止盈后切换防守","+20%","-5%","底仓"),
]
for i,(n,c,a,entry,target,stop,profit) in enumerate(plans):
    for j,v in enumerate([n,c,a,entry,target,stop,profit]):
        pt.cell(i+1,j).text = v

p = doc.add_paragraph()
r = p.add_run('\n总投入：6万 → 目标：+2万（+33%）→ 最大风险：-4200（-7%）')
r.bold = True

# 时间表
doc.add_heading('五、周执行时间表', level=1)
for day,action in [
    ("周一开盘","昆仑万维4万+东方财富2万建仓"),
    ("周一~周三","持仓，昆仑+15%出一半锁利"),
    ("周三收盘","若未盈利减半仓"),
    ("周四","达+30%目标全部止盈"),
    ("周五收盘","无论盈亏清仓"),
]:
    p = doc.add_paragraph()
    r = p.add_run(f'{day}：'); r.bold = True
    p.add_run(action)

# 风控
doc.add_heading('六、风险控制', level=1)
for risk in [
    "昆仑万维单票4万，跌停一天亏8000，必须-7%止损",
    "RSI 70.3偏热，高开+5%以上不追",
    "设条件单：昆仑万维跌-7%自动止损",
    "严格执行计划，不因【再等等】而放松纪律",
]:
    doc.add_paragraph(risk, style='List Bullet')

# 八字逻辑
doc.add_heading('七、八字选股逻辑', level=1)
doc.add_paragraph(
    '命主丁火日主，生于戌月土旺，年时双戌，土极旺而晦火。'
    '全局缺金少水，用神取金泄土、取水调候。'
    '选股以行业属金（证券/有色/保险）、属水（AI/软件/白酒/航运/半导体）为首选。'
    '代码尾数1/6/9（壬水/癸水）、7/8（庚金/辛金）加分。'
    '忌火土行业（新能源/光伏/医药/煤炭/建材）。'
)

# 免责
doc.add_heading('免责声明', level=3)
doc.add_paragraph(
    '本报告基于公开数据与八字命理模型生成，仅供学术研究参考，不构成投资建议。'
    '股市有风险，投资需谨慎。一周33%目标为极高风险操作，请量力而行。'
)

path = OUT / f'股票研报_6万8万计划_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
doc.save(str(path))
print(f"✅ {path}")
for s in scored[:3]:
    print(f"  #{scored.index(s)+1} {s[0]} ({s[1]}) {s[-1]}分")
